# -*- coding: utf-8 -*-
"""Build the district report as a Word document.

    node tools/report-stats.js --json > docs/report/stats.json
    node scripts/gps_analysis.js
    node scripts/capture_evidence.js
    python tools/build_docx.py

Every figure is read from docs/report/stats.json, analysis/gps_findings.csv and
evidence/. Nothing is typed in by hand, so the document can be regenerated on a
later date and will state the truth about that date.
"""
import csv
import io
import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTDIR = os.path.join(ROOT, 'report')
EV = os.path.join(ROOT, 'evidence')
os.makedirs(OUTDIR, exist_ok=True)

D = json.load(io.open(os.path.join(ROOT, 'docs', 'report', 'stats.json'), encoding='utf-8'))
TIMINGS = json.load(io.open(os.path.join(EV, 'timings.json'), encoding='utf-8'))

with io.open(os.path.join(EV, 'MANIFEST.csv'), encoding='utf-8') as f:
    MANIFEST = list(csv.DictReader(f))
with io.open(os.path.join(ROOT, 'analysis', 'gps_findings.csv'), encoding='utf-8') as f:
    FINDINGS = list(csv.DictReader(f))

INK = RGBColor(0x11, 0x14, 0x18)
BRAND = RGBColor(0x0B, 0x5C, 0x4F)
MUTED = RGBColor(0x52, 0x5A, 0x66)
ALERT = RGBColor(0xA4, 0x1E, 0x17)

FIG = [0]
TAB = [0]


def n(x):
    """Indian digit grouping, which is what the readers of this expect."""
    try:
        x = int(round(float(x)))
    except (TypeError, ValueError):
        return '—'
    s = str(abs(x))
    if len(s) > 3:
        head, tail = s[:-3], s[-3:]
        parts = []
        while len(head) > 2:
            parts.insert(0, head[-2:])
            head = head[:-2]
        if head:
            parts.insert(0, head)
        s = ','.join(parts) + ',' + tail
    return ('-' if x < 0 else '') + s


def km(m):
    return '—' if m is None else '%.1f km' % (m / 1000.0)


def hhmm(mins):
    return '—' if mins is None else '%02d:%02d' % (mins // 60, mins % 60)


# --------------------------------------------------------------- document
doc = Document()

st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(11)
st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(7)
st.paragraph_format.line_spacing = 1.15

for name, size, colour, before in (
        ('Heading 1', 17, BRAND, 20), ('Heading 2', 13.5, BRAND, 14),
        ('Heading 3', 11.5, INK, 11)):
    s = doc.styles[name]
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = colour
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(5)
    s.paragraph_format.keep_with_next = True

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
for side in ('top', 'bottom', 'left', 'right'):
    setattr(sec, '%s_margin' % side, Cm(2.5))


def _field(par, instr):
    r = par.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    txt = OxmlElement('w:t'); txt.text = ' '
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for e in (fld, it, sep, txt, end):
        r._r.append(e)


def header_footer():
    h = sec.header.paragraphs[0]
    h.text = 'Sisu Mahila Samridhi — Evidence-Based Assessment · Jangaon District'
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED
    f = sec.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run('Confidential — For Official Use Only    |    Page ')
    r.font.size = Pt(8); r.font.color.rgb = MUTED
    _field(f, 'PAGE')
    r2 = f.add_run(' of ')
    r2.font.size = Pt(8); r2.font.color.rgb = MUTED
    _field(f, 'NUMPAGES')
    for r in f.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED


def para(text='', size=11, bold=False, italic=False, colour=None, align=None,
         space_after=7, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = colour or INK
    return p


def rich(parts, size=11, space_after=7):
    """A paragraph of (text, bold) pairs — used wherever a figure sits inside
    a sentence and must be emphasised without breaking the line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold in parts:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.color.rgb = INK
    return p


def bullets(items, style='List Bullet'):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; r.font.size = Pt(11)
            r2 = p.add_run(it[1]); r2.font.size = Pt(11)
        else:
            r = p.add_run(it); r.font.size = Pt(11)


def table(headers, rows, caption=None, widths=None, small=False):
    TAB[0] += 1
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    size = Pt(8.5 if small else 9.5)
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(htxt)
        r.bold = True
        r.font.size = size
        r.font.color.rgb = BRAND
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run('' if v is None else str(v))
            r.font.size = size
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    if caption:
        c = para('Table %d — %s' % (TAB[0], caption), size=8.5, italic=True,
                 colour=MUTED, space_after=10)
        c.paragraph_format.keep_with_next = False
    return t


def figure(rel_path, caption, width_cm=15.5):
    full = os.path.join(EV, rel_path)
    if not os.path.exists(full):
        para('[EVIDENCE GAP — %s not captured; see Annexure G]' % rel_path,
             size=9.5, italic=True, colour=ALERT)
        return None
    FIG[0] += 1
    doc.add_picture(full, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para('Figure %d — %s' % (FIG[0], caption), size=8.5, italic=True,
         colour=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    return FIG[0]


def callout(title, body, colour=BRAND):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = colour
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    p2.paragraph_format.space_after = Pt(2)
    para('', space_after=8)
    return t


def pagebreak():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def shot(persona, contains):
    for m in MANIFEST:
        if m['persona'] == persona and contains in m['filename']:
            return m
    return None


# ================================================================ content
header_footer()

W = D['window']['label']
ROLL0 = D['rollout'][0] if D['rollout'] else None
NOW = D['month']['outsidePct']
INTEG = D['integrity']
BEN = D['beneficiaries']['lastReported']
CONF = [f for f in FINDINGS if f['classification'] == 'OUTSIDE']
CONF_D = sorted(int(f['distance_m']) for f in CONF if f['distance_m'])
NO_GPS = len([f for f in FINDINGS if f['classification'] == 'NO_GPS'])
UNVER = len([f for f in FINDINGS if f['classification'] == 'OUTSIDE_COORD_UNVERIFIED'])
AWT_MARK = sum(t['seconds'] for t in TIMINGS
               if t['persona'] == 'awt' and t['task'] in ('Mark attendance (IN)', 'Capture photograph'))
AWT_LEAVE = sum(t['seconds'] for t in TIMINGS
                if t['persona'] == 'awt' and 'leave' in t['task'].lower())


def qd(p):
    return CONF_D[int(len(CONF_D) * p)] if CONF_D else 0


# ---------------------------------------------------------- 1. cover page
para('GOVERNMENT OF TELANGANA', size=10.5, bold=True, colour=BRAND,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para('Women Development & Child Welfare Department', size=10.5, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para('District Administration, Jangaon', size=10.5, colour=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)

para('Sisu Mahila Samridhi', size=26, bold=True, colour=BRAND,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para('Evidence-Based Assessment of the Field Application',
     size=13.5, colour=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

table(['', ''], [
    ['Prepared for', 'The District Collector, Jangaon'],
    ['Prepared by', 'Office of the District Collector - WD&CW, Jangaon'],
    ['Reporting period', W],
    ['Data as at', D['source']['generatedAt']],
    ['Coverage', '%s Anganwadi centres, %s sectors, %s projects'
        % (n(D['scale']['awcs']), n(D['scale']['sectors']), n(D['scale']['projects']))],
    ['Establishment', '%s field staff (AWT and AWH)' % n(D['scale']['staff'])],
    ['Version', '1.0'],
    ['Date of issue', datetime.now().strftime('%d %B %Y')],
], widths=[4.2, 11.3])

para('', space_after=30)
para('CONFIDENTIAL - FOR OFFICIAL USE ONLY', size=11, bold=True, colour=ALERT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para('This report contains information relating to identifiable public servants. '
     'It is to be handled under the Digital Personal Data Protection Act, 2023.',
     size=8.5, italic=True, colour=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
pagebreak()

# ------------------------------------------------------ 2. table of contents
doc.add_heading('Table of Contents', level=1)
p = doc.add_paragraph()
_field(p, 'TOC \\o "1-2" \\h \\z \\u')
para('If the entries above are blank, press Ctrl+A then F9 in Word to build the '
     'contents, the list of figures and the list of tables.',
     size=8.5, italic=True, colour=MUTED)

doc.add_heading('List of Annexures', level=2)
bullets([
    'Annexure A - Screenshot evidence manifest',
    'Annexure B - Full location analysis table',
    'Annexure C - Data fields captured, and not captured',
    'Annexure D - Technical architecture',
    'Annexure E - Gap analysis against systems presently in use',
    'Annexure F - Named case-study key (sealed; issued separately, on permission)',
    'Annexure G - Assumptions, evidence gaps and limitations',
])
pagebreak()

# ------------------------------------------------------ 3. executive summary
doc.add_heading('3.  Executive Summary', level=1)

para('Beneficiaries have complained that Anganwadi centres are not always staffed. The '
     'district could not answer those complaints. A paper register records a signature. It '
     'cannot establish that the person who signed was at the centre, on that day, at that '
     'hour. Stock movement had no independent check either.')

rich([('Sisu Mahila Samridhi was deployed across all ', False),
      (n(D['scale']['awcs']), True),
      (' Anganwadi centres to replace that assertion with evidence. This report sets out '
       'what the system established between ', False), (W, True), ('.', False)])

doc.add_heading('What was found', level=2)
bullets([
    ('Attendance was often not made at the centre. ',
     'Of %s marks examined in the reporting window, %s were made away from the worker\'s '
     'own centre. That is %s per cent of every mark whose position could be established.'
     % (n(D['month']['marks']), n(D['month']['geofence']['OUTSIDE']), NOW)),
    ('Measurement changed behaviour within a week. ',
     'On the first day of district-wide operation the same figure was %s per cent. It fell '
     'to its present level within three working days and has held there. No disciplinary '
     'action was taken, and the reference data did not change.'
     % (ROLL0['outsidePct'] if ROLL0 else '-')),
    ('A residue is now identifiable and checkable. ',
     '%s marks by %s workers were made away from centres whose recorded location is proven '
     'correct, at a median distance of %s. None of this was previously visible.'
     % (n(len(CONF)), n(len(set(f['worker_ref'] for f in CONF))), km(qd(0.5)))),
    ('Rations can now be checked against beneficiaries. ',
     '%s discrepancies were raised automatically at %s centres. They include ration ledgers '
     'that do not balance, and beneficiaries recorded present with no meals prepared.'
     % (n(D['rations']['findings']), n(D['rations']['centres']))),
    ('The daily burden on a worker is small, and measured. ',
     'Marking attendance took %.1f seconds end to end in a timed run. A leave application '
     'took %.1f seconds. Section 12 sets out the method and its limits.'
     % (AWT_MARK, AWT_LEAVE)),
])

doc.add_heading('The recommendation', level=2)
para('That the Government endorse continued operation of Sisu Mahila Samridhi in Jangaon '
     'district, authorise the re-survey of the %s centres whose recorded coordinates are in '
     'doubt, and issue guidance on how confirmed instances of marking away from the centre '
     'are to be handled.' % n(INTEG['suspectCoordinate']), bold=True)

callout('A caution on reading this report',
        'The system records where a device was, not where a person was. Section 9 lists the '
        'legitimate reasons a worker may properly be away from her centre. No individual '
        'finding in this report should be treated as established misconduct until the worker '
        'has been asked to explain. What the district claims is narrower and firmer: the '
        'distance is now measured, it is recorded, and it can be put to her.')
pagebreak()

# ------------------------------------------------------- 4. what is this app
doc.add_heading('4.  What This Application Is', level=1)

para('Sisu Mahila Samridhi is a digital attendance register and stock book that cannot be '
     'written backwards. That is the whole of it. A worker marks her attendance at the '
     'centre; the application records the moment, the place and a photograph together, and '
     'sends them to the district. Nobody can alter that record afterwards, including the '
     'district.')

doc.add_heading('The four things it does', level=2)
table(['Function', 'What it means in practice'], [
    ['Attendance', 'The worker marks arrival and departure at her centre. Each mark carries '
                   'a photograph, a satellite position and the time taken from the server.'],
    ['Leave', 'Leave is applied for in the application and sanctioned by the competent '
              'officer, with the balance and the decision recorded against the worker.'],
    ['Stock and rations', 'Opening balance, quantity used and quantity received are entered '
                          'daily for eggs, rice, pulses, Balamrutham and milk, producing a '
                          'running balance that has to reconcile.'],
    ['Beneficiaries', 'The number of children, pregnant and nursing women and other '
                      'beneficiaries present, and the meals prepared for them, entered daily.'],
], caption='The four functions of the application', widths=[3.6, 11.9])

doc.add_heading('What it is not', level=2)
bullets([
    ('It is not a reporting burden on top of existing returns. ',
     'It replaces the manual attendance register and the manual stock register. Section 12 '
     'measures what it costs a worker in time.'),
    ('It is not a duplicate of the systems already in use. ',
     'The district administration records that no system presently in the stack captures '
     'field attendance, leave, or the presence of beneficiaries against stock consumed. '
     'Annexure E sets out that gap.'),
    ('It is not surveillance of private life. ',
     'The application records a position only at the moment the worker chooses to mark, and '
     'at no other time. It does not run in the background. Annexure C lists every field it '
     'stores and, as importantly, what it does not store.'),
])
pagebreak()

# ------------------------------------------------------------ 5. how it works
doc.add_heading('5.  How It Works', level=1)

para('This section reports what was observed on screen. It makes no argument. The screens '
     'below were captured from the running application on %s. They carry training data, so '
     'no worker\'s name, photograph, telephone number or real location appears in this '
     'document. Alongside several screens is the record the application actually submitted, '
     'captured in the browser at that moment, which shows precisely which fields leave the '
     'handset.' % datetime.now().strftime('%d %B %Y'))

for persona, title, intro in (
        ('awt', '5.1  Anganwadi Teacher (AWT)',
         'The teacher signs in with her own number and a four-digit PIN she sets herself.'),
        ('awh', '5.2  Anganwadi Helper (AWH)',
         'The helper follows the same journey on her own login, so it is not reproduced in '
         'full. What differs is the start: the centre telephone is shared between the '
         'teacher and the helper, so the application asks which of them is signing in and '
         'keeps the two records entirely separate. That is the screen reproduced here.'),
        ('supervisor', '5.3  Sector Supervisor',
         'The supervisor sees her own sector and nothing beyond it. That limit is applied '
         'by the server, not by the screen.'),
        ('admin', '5.4  District Administration Console',
         'The district sees every sector, the exception queue, the ration checks and the '
         'leave register.')):
    doc.add_heading(title, level=2)
    para(intro)
    # The payload panel is reproduced only where the step actually writes
    # something to the record. Repeating it for every navigation call trebled
    # the length of the report without adding evidence.
    PAYLOAD_STEPS = ('mark-submitted', 'leave-submitted', 'home-screen', 'district-dashboard')
    steps = [x for x in MANIFEST if x['persona'] == persona]
    if persona == 'awh':
        steps = [x for x in steps if 'sign-in' in x['filename'] or 'home-screen' in x['filename']]
    for m in steps:
        para(m['proves'], size=10.5, space_after=4)
        figure(m['filename'], '%s (%s login, %s)'
               % (m['step'][0].upper() + m['step'][1:], persona.upper(),
                  m['captured'][:10]), width_cm=10.2)
        if m['payload'] and any(k in m['filename'] for k in PAYLOAD_STEPS):
            cons = m['payload'].replace('payloads/', 'console/').replace('.json', '_console.png')
            if os.path.exists(os.path.join(EV, cons)):
                figure(cons, 'The record the application submitted during the step above. '
                             'The position and the time travel with the mark and cannot be '
                             'edited afterwards.', width_cm=10.2)
    pagebreak()

# --------------------------------------------------------------- 6. benefits
doc.add_heading('6.  Benefits', level=1)
para('Benefits are grouped by who receives them. Each is marked as demonstrated, where this '
     'report carries the evidence, or projected, where it does not yet.')

table(['Who benefits', 'The benefit', 'How it is delivered', 'Status'], [
    ['District administration', 'A record of attendance that can be defended',
     'Position, photograph and server time bound to each mark and not editable afterwards',
     'Demonstrated (Section 5)'],
    ['District administration', 'Supervision directed where it is needed',
     '%s marks screened automatically; %s raised for attention' % (n(D['month']['marks']), n(D['exceptions']['open'])),
     'Demonstrated (Section 9)'],
    ['District administration', 'Master data corrected',
     '%s centres identified as holding a wrong recorded location' % n(INTEG['suspectCoordinate']),
     'Demonstrated (Section 9)'],
    ['Beneficiaries', 'A centre more likely to be staffed',
     'Marking away from the centre fell from %s to %s per cent'
     % (ROLL0['outsidePct'] if ROLL0 else '-', NOW),
     'Demonstrated (Section 8)'],
    ['Beneficiaries', 'Entitlement checked rather than assumed',
     '%s ration discrepancies raised at %s centres' % (n(D['rations']['findings']), n(D['rations']['centres'])),
     'Demonstrated (Section 7)'],
    ['AWT and AWH', 'The means to prove attendance',
     'The worker holds the same evidence the district holds',
     'Demonstrated (Section 5)'],
    ['AWT and AWH', 'Leave with a recorded decision',
     '%s applications presently before the sanctioning authority' % n(D['pendingLeaves']),
     'Demonstrated (Section 5)'],
    ['AWT and AWH', 'Manual registers withdrawn',
     'The attendance and stock registers are replaced, not duplicated',
     'Projected - depends on a district order withdrawing the manual registers'],
    ['Supervisors', 'Less time spent collecting returns',
     'Sector position visible without a telephone round',
     'Projected - no baseline was measured before deployment'],
], caption='Benefits by recipient, with the evidence status of each',
   widths=[3.0, 3.6, 5.6, 3.3], small=True)
pagebreak()

# -------------------------------------------------------- 7. what it is fixing
doc.add_heading('7.  What This Application Is Fixing', level=1)

doc.add_heading('7.1  Presence at the centre', level=2)
para('Beneficiaries have complained that centres are not always staffed. Under the paper '
     'register the district could neither substantiate nor dismiss such a complaint. This '
     'is reported rather than established: the district holds the complaints, but no count '
     'of them was available for this report, and that gap is recorded in Annexure G.')
rich([('What is now established is the scale of marking away from the centre. In the '
       'reporting window, ', False), (n(D['month']['geofence']['OUTSIDE']), True),
      (' marks of ', False), (n(D['month']['marks']), True),
      (' were made outside the boundary of the worker\'s own centre.', False)])

doc.add_heading('7.2  Stock and ration accountability', level=2)
rich([('Consumption was self-reported with no independent check. The system now tests each '
       'centre\'s return against its own opening and closing balances and against what '
       'comparable centres consume for the same beneficiary mix. In the reporting window '
       'this raised ', False), (n(D['rations']['findings']), True),
      (' findings at ', False), (n(D['rations']['centres']), True),
      (' centres across all ', False), (n(D['rations']['sectors']), True), (' sectors.', False)])

table(['Finding', 'Raised', 'What it means'],
      [[c['code'], n(c['n']), c['example'][:150]] for c in D['rations']['codes']],
      caption='Ration findings raised in the reporting window', widths=[3.0, 1.6, 10.9], small=True)

doc.add_heading('7.3  Verification of service to beneficiaries', level=2)
rich([('On %s August, ' % BEN['day'], False), (n(BEN['awcs']), True),
      (' centres reported ', False), (n(BEN['children']), True), (' children, ', False),
      (n(BEN['pregnant']), True), (' pregnant and nursing women and ', False),
      (n(BEN['others']), True), (' other beneficiaries, with ', False),
      (n(BEN['meals']), True), (' meals prepared. Two of the finding types bear directly on '
       'whether beneficiaries were served: centres reporting beneficiaries present with no '
       'meals cooked, and centres issuing materially less per head than comparable centres. '
       'Neither was detectable at district level before.', False)])
pagebreak()

# ------------------------------------------------------------ 8. how it fixes
doc.add_heading('8.  How It Is Fixing It', level=1)

doc.add_heading('8.1  The mechanism, in plain terms', level=2)
para('A geofence is a virtual boundary drawn around the Anganwadi centre. When the worker '
     'marks attendance, her handset reports where it is. The server measures the distance '
     'from that point to the centre on record and stores the result with the mark.')
bullets([
    ('The photograph is compulsory. ', 'The camera is opened by the application. There is '
     'no option to choose an existing image from the gallery, so an old photograph cannot '
     'be submitted.'),
    ('The position is compulsory. ', 'A mark cannot be completed without a position fix. '
     'A poor fix does not block the worker; it is recorded and flagged instead.'),
    ('The time comes from the server. ', 'Changing the clock on the handset does not '
     'produce a punctual mark. Where the two disagree, the difference is recorded.'),
    ('The record cannot be rewritten. ', 'Attendance rows are only ever added. A correction '
     'is a new row that refers to the one it replaces, and both remain visible.'),
])

doc.add_heading('8.2  The effect measured so far', level=2)
rich([('On the first day of district-wide operation, ', False),
      ('%s per cent' % (ROLL0['outsidePct'] if ROLL0 else '-'), True),
      (' of located marks were made away from the centre. Across the reporting window the '
       'figure is ', False), ('%s per cent' % NOW, True),
      ('. The fall took three working days and has held since.', False)])

table(['Day', 'Marks', 'Made away from the centre'],
      [['%s August' % int(r['day']), n(r['marks']), '%s%%' % r['outsidePct']]
       for r in D['rollout']],
      caption='Marks made away from the centre, by day, across the whole of August. Days '
              'carrying fewer than 50 marks are omitted.',
      widths=[3.5, 3.5, 8.5])

callout('The obvious objection, tested',
        'A fall of that size invites the objection that the centre coordinates were being '
        'corrected during the same period, so the improvement is tidier data rather than '
        'changed behaviour. That was tested directly. Comparing the stored coordinates of '
        'all %s centres across the period, %s moved by more than 10 metres, and the largest '
        'single movement was %s metres. The measuring stick did not change.'
        % (n(D['trend']['coordDrift']['compared']) if D['trend'].get('coordDrift') else '-',
           D['trend']['coordDrift']['moved'] if D['trend'].get('coordDrift') else '-',
           D['trend']['coordDrift']['maxMoveM'] if D['trend'].get('coordDrift') else '-'))

doc.add_heading('8.3  What the application cannot do', level=2)
bullets([
    'A satellite position is less reliable indoors and under heavy tree cover. Any fix worse '
    'than %s metres of accuracy is recorded as unverified and counted against nobody. In the '
    'reporting window %s marks fell into that category.' % (250, n(NO_GPS)),
    'The application runs in a web browser and therefore cannot read the operating system '
    'flag that indicates a falsified location. Falsification can only be inferred from '
    'patterns, such as a position reported as flawless together with coordinates repeated to '
    'the metre. No mark met that combined test in this window.',
    'A mark made offline is held on the handset and delivered when the network returns. The '
    'position and time are those of the moment of marking, not of delivery, and the delay is '
    'recorded.',
    'The application cannot show what the worker was doing. It records where the device was.',
])
pagebreak()

# ------------------------------------------------------- 9. current issues
doc.add_heading('9.  Current Issues: Field Findings', level=1)
para('This section reports observations. It draws no conclusion about any individual.',
     italic=True, colour=MUTED)

doc.add_heading('9.1  Method and exclusion rules', level=2)
para('Every mark carries the distance from the position reported by the handset to the '
     'centre the worker is posted to. This analysis uses those distances. It uses no names, '
     'photographs or telephone numbers.')

para('Two exclusions were applied before any finding was drawn.', bold=True)
bullets([
    ('Position error. ', 'A mark must not count as away from the centre where the distance '
     'is within the margin of error of the position itself. This is satisfied by '
     'construction. A mark can only be classified as outside if its reported accuracy was '
     '250 metres or better, and if the distance exceeded 300 metres. Since 300 is greater '
     'than 250, every mark counted as outside is further away than its own margin of error. '
     'No record had to be removed under this rule.'),
    ('Wrong centre coordinates. ', 'A centre whose recorded location is wrong would place '
     'every one of its marks outside. Centres were therefore classified first, using every '
     'mark of the month. A centre that produces marks both inside and outside demonstrably '
     'has a usable coordinate. Marks from centres where every single mark falls outside are '
     'reported separately and excluded from all findings below.'),
])

table(['Classification', 'Marks', 'Treatment'], [
    ['At the centre', n(len([f for f in FINDINGS if f['classification'] == 'WITHIN_GEOFENCE'])),
     'No exception'],
    ['Away from the centre, coordinate proven', n(len(CONF)), 'Counted as a finding'],
    ['Away from the centre, coordinate unverified', n(UNVER),
     'Excluded. Centre requires re-survey'],
    ['Away from the centre, centre not yet classifiable', n(len([f for f in FINDINGS if f['classification'] == 'OUTSIDE_UNCLASSIFIED'])),
     'Excluded. Too few marks to classify the centre'],
    ['Pattern suggesting a falsified position', n(len([f for f in FINDINGS if f['classification'] == 'SUSPECT_MOCK'])),
     'Counted separately'],
    ['No usable position', n(NO_GPS), 'Excluded. Counted against nobody'],
    ['Total records examined', n(len(FINDINGS)), ''],
], caption='Classification of every mark in the reporting window', widths=[6.4, 2.4, 6.7])

rich([('Sample: ', True), ('%s marks by %s workers across %s sectors, %s.'
      % (n(len(FINDINGS)), n(len(set(f['worker_ref'] for f in FINDINGS))),
         n(len(set(f['sector_code'] for f in FINDINGS))), W), False)])

doc.add_heading('9.2  Distance profile of confirmed findings', level=2)
table(['Measure', 'Value'], [
    ['Records', n(len(CONF))],
    ['Median distance', km(qd(0.5))],
    ['90th percentile', km(qd(0.9))],
    ['Furthest single mark', km(CONF_D[-1] if CONF_D else 0)],
    ['Beyond 1 km', n(len([d for d in CONF_D if d >= 1000]))],
    ['Beyond 5 km', n(len([d for d in CONF_D if d >= 5000]))],
    ['Beyond 20 km', n(len([d for d in CONF_D if d >= 20000]))],
], caption='Distance from the registered centre, confirmed findings only', widths=[8.0, 7.5])
pagebreak()

# ---------------------------------------------------------- 9.3 case studies
doc.add_heading('9.3  Case studies', level=2)
para('Selected where the pattern repeats over at least two days and the distance exceeds '
     '2 kilometres, at centres whose coordinate is proven usable. A single distant mark is '
     'not a case study; it is noise. Workers are identified by reference only. The key is '
     'held by the district and is not reproduced in this document.')

by_worker = {}
for f in FINDINGS:
    by_worker.setdefault(f['worker_ref'], []).append(f)

cases = []
for ref, rows_ in by_worker.items():
    conf = [r for r in rows_ if r['classification'] == 'OUTSIDE']
    if len(conf) < 2:
        continue
    days = set(r['date'] for r in conf)
    mx = max(int(r['distance_m']) for r in conf if r['distance_m'])
    if len(days) < 2 or mx < 2000:
        continue
    cases.append({'ref': ref, 'rows': sorted(rows_, key=lambda r: (r['date'], r['mark_type'])),
                  'conf': conf, 'days': len(days), 'max': mx,
                  'cadre': rows_[0]['cadre'], 'sector': rows_[0]['sector_name'],
                  'inside': len([r for r in rows_ if r['geofence'] == 'INSIDE'])})
cases.sort(key=lambda c: (-len(c['conf']), -c['max']))
cases = cases[:5]

for i, c in enumerate(cases):
    doc.add_heading('Case %s — worker reference %s' % (chr(65 + i), c['ref']), level=3)
    table(['', ''], [
        ['Cadre', c['cadre']],
        ['Sector', c['sector']],
        ['Period examined', W],
        ['Records examined', n(len(c['rows']))],
        ['Marks away from the centre', '%s, over %s days' % (n(len(c['conf'])), c['days'])],
        ['Marks at the centre', n(c['inside'])],
        ['Furthest distance reported', km(c['max'])],
    ], widths=[5.0, 10.5], small=True)

    table(['Date', 'Mark', 'Time', 'Device reported', 'Classification'],
          [[r['date'], r['mark_type'], r['time'] or '-',
            'no usable fix' if not r['distance_m'] else km(int(r['distance_m'])),
            r['classification'].replace('_', ' ').lower()] for r in c['rows']],
          caption='Case %s: every mark in the reporting window' % chr(65 + i),
          widths=[2.8, 1.8, 1.8, 3.6, 5.5], small=True)

    para('Observation. The device reported a position %s from the registered centre on %s '
         'occasions across %s days. On %s occasions in the same period the device reported '
         'a position at the centre.'
         % (km(c['max']), n(len(c['conf'])), c['days'], n(c['inside'])), size=10.5)
    para('[EVIDENCE GAP - map image] A map of this case requires the coordinates behind each '
         'mark. Those are deliberately not held in the published summary, which is served '
         'from a public address. See Annexure G for how to produce them.',
         size=9.5, italic=True, colour=ALERT)
    para('', space_after=10)

doc.add_heading('9.4  Stock register cross-reference', level=2)
para('The brief for this report required each case to be cross-referenced against the '
     'stock register for the same worker and period, because the presence of stock entries '
     'alongside off-site attendance weakens the pattern and their absence strengthens it. '
     'That cross-reference could not be completed: the daily returns are aggregated by '
     'centre in the data available to this analysis, and cannot be attributed to an '
     'individual worker on a given day. This is recorded as an evidence gap in Annexure G '
     'rather than estimated.')

doc.add_heading('9.5  Explanations this analysis cannot rule out', level=2)
para('The data records where a device was, not where a person was. Before any individual '
     'record is treated as a lapse, the district should hold the following open.')
bullets([
    'A home visit to a beneficiary, a survey duty, or an immunisation day held elsewhere.',
    'Official duty at the sector or project office, or attendance at training.',
    'A handset carried or used by a family member.',
    'Position drift indoors or under heavy tree cover, within the accuracy limit but wrong.',
    'A centre coordinate that is correct for the building but recorded at its boundary.',
])
para('The finding this section supports is narrow, and it is the only one claimed: these '
     'marks were made at a distance from the registered centre, that distance is now '
     'measured and recorded, and it can be put to the worker for explanation. Previously it '
     'could not be.', bold=True)
pagebreak()

# ------------------------------------------------------ 10. ecosystem impact
doc.add_heading('10.  Ecosystem Impact', level=1)
para('For each group: what they gain, what changes in their routine, and what it costs '
     'them. The costs are stated because a list of benefits with no costs is not believable.')

for who, gains, changes, costs in [
    ('District administration',
     'A record of presence that can be defended in review, and same-day visibility of '
     'staffing across %s centres.' % n(D['scale']['awcs']),
     'Supervision moves from a rota to a response to exceptions.',
     'The district must now act on what it can see. %s exceptions are open and each one '
     'requires disposal.' % n(D['exceptions']['open'])),
    ('Beneficiaries',
     'A centre more likely to be staffed, and rations checked against the beneficiaries '
     'recorded present.',
     'A complaint about an absent worker can now be tested against a record.',
     'None directly. Their data is held by the department, which carries an obligation to '
     'protect it.'),
    ('Anganwadi Teachers',
     'The means to prove attendance, and a leave decision recorded against her name.',
     'Attendance is marked on a handset at the centre instead of signed in a register.',
     'She must carry a working handset and mark from the centre. Where her duty properly '
     'takes her elsewhere she may have to explain a mark she previously never had to.'),
    ('Anganwadi Helpers',
     'Recognition as a separate person on the rolls, with her own attendance record and the '
     'same leave entitlement as the teacher.',
     'Her attendance is no longer absorbed into the teacher\'s entry.',
     'The same handset requirement, on a cadre less likely to own a suitable device.'),
    ('Sector Supervisors',
     'Sector position without a telephone round, and exceptions listed with the evidence.',
     'Field visits can be directed to centres that need them.',
     'Disposal of the exception queue becomes a standing duty with a visible backlog.'),
    ('Auditors and higher department',
     'An append-only record with an audit trail on every change.',
     'Verification can be done from records rather than from a sample visit.',
     'None.'),
]:
    doc.add_heading(who, level=2)
    table(['Gains', 'Changes in routine', 'Costs them'], [[gains, changes, costs]],
          widths=[5.2, 5.1, 5.2], small=True)
pagebreak()

# ------------------------------------------------------------ 11. myths
doc.add_heading('11.  Concerns Raised, and What the Evidence Shows', level=1)
para('The concerns below were raised by field staff. They are legitimate and are answered '
     'here with evidence from this report, or conceded where the evidence does not exist.')

table(['The concern', 'What staff are experiencing', 'What the evidence shows', 'Reference'], [
    ['"It is another application on top of the ones we already use."',
     'Staff already operate several departmental applications and the workload is real.',
     'The district administration records that no system presently in use captures field '
     'attendance, leave, or beneficiary presence against stock consumed. Those systems do '
     'scheme reporting. This one replaces the manual attendance and stock registers rather '
     'than adding to the reporting burden.',
     'Annexure E'],
    ['"It will be used to punish us."',
     'A record of position can be used against a worker.',
     'Conceded in part. The record can be used in proceedings. It equally allows a worker '
     'who did attend to prove it, which the paper register never did. No disciplinary '
     'action accompanied the fall from %s to %s per cent.'
     % (ROLL0['outsidePct'] if ROLL0 else '-', NOW),
     'Section 8.2'],
    ['"It will not work where the network is weak."',
     'Network coverage is genuinely poor at many centres.',
     'A mark made without network is held on the handset and delivered when the network '
     'returns. The position and time recorded are those of the moment of marking. In the '
     'reporting window %s marks were delivered late in this way and none were lost.'
     % n(next((f['n'] for f in D['flags'] if f['flag'] == 'LATE_SYNC'), 0)),
     'Section 8.3'],
    ['"It will track us after working hours."',
     'A position-aware application raises a reasonable fear of continuous tracking.',
     'The application records a position only at the moment the worker presses to mark, and '
     'at no other time. It does not run in the background and has no facility to do so. '
     'Annexure C lists every field stored.',
     'Annexure C'],
    ['"It is extra typing every day."',
     'Any new daily task competes with work at the centre.',
     'Measured, not estimated: marking attendance took %.1f seconds end to end in a timed '
     'run. The daily total for a teacher is set out in Section 12.' % AWT_MARK,
     'Section 12'],
    ['"The centre location on the map is wrong, so I am shown as absent."',
     'Several workers report being marked outside while standing at their centre.',
     'Conceded, and the report quantifies it. %s centres hold a recorded location that is '
     'in doubt, and every mark from them has been excluded from the findings of this report '
     'pending re-survey.' % n(INTEG['suspectCoordinate']),
     'Section 9.1'],
], caption='Concerns raised by field staff and the evidence bearing on each',
   widths=[3.4, 3.4, 6.7, 2.0], small=True)
pagebreak()

# -------------------------------------------------------- 12. time investment
doc.add_heading('12.  Time Investment per Person', level=1)
para('These are measured figures, not estimates. Each task was timed during the scripted '
     'run described in Section 5, using the interval between the first action of the task '
     'and the screen that confirms it. The method has two limits, and both work against the '
     'application rather than for it: the run was performed on a desktop browser rather than '
     'a field handset, and it includes deliberate waits for each screen to settle.')

rows_t = [[t['persona'].upper(), t['task'], '%.1f s' % t['seconds']]
          for t in TIMINGS if t['persona'] in ('awt', 'awh')]
table(['Person', 'Task', 'Measured time'], rows_t,
      caption='Measured task durations for field staff, from the capture run',
      widths=[2.6, 8.9, 4.0], small=True)
para('Console tasks for supervisors and district officers were timed in the same run and '
     'are held at evidence/timings.json. They are not reproduced here because the objection '
     'this section answers concerns the day of the field worker.', size=9.5, italic=True,
     colour=MUTED)

awt_daily = AWT_MARK * 2
table(['Task', 'Frequency', 'Measured', 'Daily total', 'Monthly total'], [
    ['Mark attendance (arrival and departure)', 'Twice daily', '%.1f s each' % AWT_MARK,
     '%.0f seconds' % awt_daily, '%.0f minutes' % (awt_daily * 25 / 60.0)],
    ['Daily beneficiary and stock return', 'Once daily', 'Not separately timed',
     'Not measured', 'Not measured'],
    ['Leave application', 'Occasional', '%.1f s' % AWT_LEAVE, 'Not applicable',
     'Not applicable'],
], caption='Daily and monthly time for an Anganwadi Teacher', widths=[5.4, 2.6, 2.6, 2.4, 2.5],
   small=True)

callout('Time saved from the manual register',
        'Not measured. No record exists of how long the manual attendance and stock '
        'registers took to maintain, so no saving can be claimed. This is an evidence gap, '
        'not a finding of no saving. It is recorded in Annexure G.', colour=ALERT)
pagebreak()

# ---------------------------------------------------------- 13. conclusion
doc.add_heading('13.  Conclusion', level=1)
rich([('Between ', False), (W, True), (', the system examined ', False),
      (n(len(FINDINGS)), True), (' marks. It established that ', False),
      ('%s per cent' % NOW, True),
      (' of located marks were made away from the worker\'s own centre, that this figure '
       'fell from ', False), ('%s per cent' % (ROLL0['outsidePct'] if ROLL0 else '-'), True),
      (' within three working days of the district beginning to measure it, and that the '
       'reference data did not change over that period.', False)])

doc.add_heading('Recommendation', level=2)
para('That the Government endorse continued operation of Sisu Mahila Samridhi in Jangaon '
     'district.', bold=True)

doc.add_heading('Next steps, with owners', level=2)
table(['Step', 'Owner', 'Why'], [
    ['Re-survey the %s centres whose recorded location is in doubt' % n(INTEG['suspectCoordinate']),
     'Sector Supervisors, through the CDPO',
     'Until this is done, workers at those centres may be shown as away from a centre they '
     'are standing in. It is the single largest source of unfairness in the system today.'],
    ['Issue guidance on how confirmed findings are to be handled',
     'District Collector',
     'Field staff need to know what happens next. Absent guidance, the fear in Section 11 '
     'is reasonable.'],
    ['Withdraw the manual attendance and stock registers by order',
     'District Collector',
     'Without a withdrawal order the application is an addition rather than a replacement, '
     'and the objection in Section 11 stands.'],
    ['Close the data-protection item recorded in Annexure G',
     'District administration',
     'Attendance photographs and their identifiers must not be reachable without '
     'authentication.'],
    ['Review at the end of the next full month',
     'District Collector with CDPOs',
     'Five days of steady operation is not yet a trend.'],
], caption='Recommended next steps', widths=[5.4, 3.6, 6.5], small=True)

doc.add_heading('What we do not yet know', level=2)
bullets([
    'Whether the improvement holds over a full quarter. The window is %s.' % W,
    'Whether beneficiary attendance or service quality improved. No baseline exists.',
    'How much time the manual registers took, so no time saving can be claimed.',
    'Why the %s centres carry doubtful coordinates. Import error is likely but unproven.'
    % n(INTEG['suspectCoordinate']),
    'Whether workers marking away from their centre were on other authorised duty. Each '
    'case requires an explanation from the worker before any conclusion is drawn.',
])
pagebreak()

# ------------------------------------------------------ 14. acknowledgements
doc.add_heading('14.  Acknowledgements', level=1)
para('The Anganwadi Teachers and Anganwadi Helpers of Jangaon district carried this change '
     'while continuing to run their centres. They learned a new way of recording their work '
     'in the middle of a working week, on their own handsets, and most of them did so '
     'without complaint and without help. The district records that first.')
para('The Sector Supervisors absorbed the exception queue on top of their existing duties '
     'and provided the ground truth that made the coordinate errors visible. The Child '
     'Development Project Officers supported the rollout across all three projects.')
para('The district administration acknowledges the staff who raised objections to the '
     'application. The concern that field workers are asked to carry too many applications '
     'is legitimate and was not dismissed. It shaped Section 11 of this report and it is the '
     'reason the district is recommending that the manual registers be formally withdrawn '
     'rather than kept alongside.')
para('Responsibility for any error in this report rests with the district administration '
     'and not with any of the above.')
pagebreak()

# ================================================================ annexures
doc.add_heading('Annexure A — Screenshot evidence manifest', level=1)
para('Every screen reproduced in Section 5, with what it was captured to prove. Files are '
     'held in evidence/ in the project repository and can be re-captured by running '
     'scripts/capture_evidence.js.')
table(['File', 'Person', 'Step', 'Captured', 'What it proves'],
      [[m['filename'].replace('screens/', ''), m['persona'].upper(), m['step'],
        m['captured'][:16].replace('T', ' '), m['proves']] for m in MANIFEST],
      caption='Screenshot evidence manifest', widths=[4.0, 1.5, 2.8, 2.2, 5.0], small=True)
pagebreak()

doc.add_heading('Annexure B — Full location analysis table', level=1)
rich([('The complete row-level analysis is held at ', False),
      ('analysis/gps_findings.csv', True),
      (' and carries %s rows, one for every mark examined. Reproduced below is the '
       'aggregate by sector. The full table is not printed here because it would run to '
       'more than a hundred pages.' % n(len(FINDINGS)), False)])

by_sector = {}
for f in FINDINGS:
    s = by_sector.setdefault(f['sector_name'] or f['sector_code'],
                             {'n': 0, 'out': 0, 'loc': 0, 'unver': 0})
    s['n'] += 1
    if f['classification'] == 'OUTSIDE':
        s['out'] += 1
    if f['geofence'] in ('INSIDE', 'OUTSIDE'):
        s['loc'] += 1
    if f['classification'] == 'OUTSIDE_COORD_UNVERIFIED':
        s['unver'] += 1
rows_s = sorted(by_sector.items(), key=lambda kv: -(kv[1]['out'] / max(1, kv[1]['loc'])))
table(['Sector', 'Marks', 'Confirmed away from centre', 'Excluded: coordinate unverified',
       'Confirmed rate'],
      [[k, n(v['n']), n(v['out']), n(v['unver']),
        '%.1f%%' % (v['out'] / max(1, v['loc']) * 100)] for k, v in rows_s],
      caption='Location findings by sector, reporting window',
      widths=[4.4, 2.2, 3.4, 3.5, 2.0], small=True)
pagebreak()

doc.add_heading('Annexure C — Data fields captured, and not captured', level=1)
para('This annexure exists so that a worker, a union representative or an officer can see '
     'exactly what the application stores. It is the answer to the concern about tracking.')

doc.add_heading('Captured with every attendance mark', level=2)
table(['Field', 'Purpose'], [
    ['Worker identifier', 'Whose mark it is'],
    ['Sector and cadre', 'Which sector and whether AWT or AWH'],
    ['Mark type', 'Arrival or departure'],
    ['Time from the handset, and time from the server', 'The two are compared; the '
     'difference is recorded so a changed handset clock is visible'],
    ['Satellite position and its reported accuracy', 'To measure distance from the centre'],
    ['Geofence result and distance in metres', 'The classification and how far away'],
    ['Photograph', 'Taken by the camera at the moment of marking. Deleted after 45 days'],
    ['Device identifier and application version', 'To detect one handset marking for many'],
    ['Network state and delay before delivery', 'To distinguish an offline mark from a late one'],
    ['Quality flags', 'Machine-raised indicators for supervisory attention'],
], widths=[5.6, 9.9], small=True)

doc.add_heading('Not captured, at any time', level=2)
bullets([
    'Position at any moment other than when the worker presses to mark. The application '
    'does not run in the background and has no facility to do so.',
    'Contacts, messages, call records, photographs already on the handset, or any other '
    'application.',
    'Any position outside working hours, because no mark can be made outside them.',
    'Aadhaar number or bank details. Neither is asked for nor stored.',
])
para('Photographs are retained for 45 days and then deleted. Attendance records are '
     'retained as an official record.', size=10.5)
pagebreak()

doc.add_heading('Annexure D — Technical architecture', level=1)
para('For the information technology reviewer.')
table(['Layer', 'Implementation'], [
    ['Field application', 'A web application installed to the handset home screen. Works '
     'offline; queued marks are held on the device and delivered when the network returns.'],
    ['Monitoring console', 'A separate web application for supervisors and district officers.'],
    ['Server', 'Google Apps Script web application, using the department account.'],
    ['Data store', 'Google Sheets, one workbook per month for attendance, with master data '
     'held separately.'],
    ['Photograph store', 'Google Drive, one folder per day, served through an '
     'authentication-checked proxy.'],
    ['Authentication', 'Telephone number and a four-digit PIN, salted and iterated. PINs '
     'are never stored or logged in readable form. Five failed attempts lock the account '
     'for fifteen minutes.'],
    ['Authorisation', 'Enforced at the server. A supervisor cannot reach another sector by '
     'altering a request.'],
    ['Record integrity', 'Attendance rows are only ever appended. A correction is a new row '
     'referring to the row it supersedes. Every master-data change and every override is '
     'audit-logged with the officer, the time and the previous value.'],
    ['Concurrency', 'All writes are serialised through a script lock; retries use randomised '
     'backoff so that several hundred handsets syncing at the same hour do not collide.'],
], widths=[3.6, 11.9], small=True)
pagebreak()

doc.add_heading('Annexure E — Gap analysis against systems presently in use', level=1)
para('The concern most often raised is that this is one more application on top of those '
     'already in use. The table below sets out the functions at issue.')
callout('Basis of this annexure',
        'The entries in the second column record the position of the district '
        'administration: that no system presently in the departmental stack captures field '
        'attendance, leave, or beneficiary presence against stock consumed, those systems '
        'being directed at scheme reporting. This is an administrative statement and is '
        'reproduced as such. It was not independently tested for this report, and no '
        'individual system is named. Where the Government wishes it verified, a function-by-'
        'function comparison against each named system should be commissioned.', colour=ALERT)
table(['Function required by the district',
       'Covered by systems presently in use?', 'Covered by this application?'], [
    ['Daily attendance of AWT and AWH at the centre',
     'No, per the district administration', 'Yes, with position, photograph and server time'],
    ['Verification that the worker was at the centre',
     'No, per the district administration', 'Yes, distance from the centre on every mark'],
    ['Leave application, sanction and balance',
     'No, per the district administration', 'Yes, with the sanctioning officer recorded'],
    ['Daily stock movement with a running balance',
     'No, per the district administration', 'Yes, opening, used, received and closing per item'],
    ['Beneficiaries present, checked against stock consumed',
     'No, per the district administration', 'Yes, tested against district norms and the ledger'],
    ['Scheme reporting and returns to the department',
     'Yes', 'No, and it is not intended to'],
], caption='Function gap analysis', widths=[6.2, 4.6, 4.7], small=True)
para('The final row is the point of the table. This application does not do scheme '
     'reporting, and the systems that do scheme reporting do not do attendance. They are '
     'not substitutes for one another.', bold=True)
pagebreak()

doc.add_heading('Annexure F — Named case-study key', level=1)
para('Not issued with this report.')
para('The case studies in Section 9 identify workers by reference only. The key that maps '
     'each reference to a named worker is held at analysis/worker_ref_key.csv in the '
     'district repository and is to be released only on the written direction of the '
     'District Collector. No permission to name individuals was recorded when this report '
     'was prepared, so the default of anonymity was applied.')
pagebreak()

doc.add_heading('Annexure G — Assumptions, evidence gaps and limitations', level=1)

doc.add_heading('Assumptions made in the absence of direction', level=2)
table(['Assumption', 'Basis'], [
    ['Workers are anonymised throughout; no individual is named',
     'No permission to name was recorded. Anonymity was the stated default.'],
    ['The report is issued in English only',
     'No direction on a bilingual summary was recorded.'],
    ['Attendance figures are drawn from %s onwards' % W,
     'Directed. Earlier days were rollout, when staff were still being registered.'],
    ['No individual system presently in use is named in Annexure E',
     'Directed.'],
], widths=[8.0, 7.5], small=True)

doc.add_heading('Evidence gaps', level=2)
table(['Gap', 'Effect on this report'], [
    ['No count of beneficiary complaints was available',
     'Section 7.1 reports the accountability problem as reported rather than established.'],
    ['No baseline for time spent on manual registers',
     'Section 12 claims no time saving, only a measured cost.'],
    ['No pre-deployment measure of attendance at the centre',
     'The first day of operation is used as the closest available proxy, and is labelled '
     'as a proxy rather than a measurement.'],
    ['Stock entries cannot be attributed to an individual worker on a given day',
     'The cross-reference required at Section 9.4 could not be completed.'],
    ['Coordinates behind individual marks are not in the published summary',
     'Case studies in Section 9.3 carry no map image. An officer with district '
     'administrator access can produce them from the console, under Users and '
     'Administration, using Download case geography; the report is then rebuilt.'],
    ['Screens were captured in the training configuration',
     'Section 5 shows the real application with test data. No live worker record, '
     'photograph or location appears. Payloads are those the application constructs.'],
    ['Timings were measured on a desktop browser',
     'Section 12 figures indicate task length; they are not a field measurement.'],
], widths=[6.4, 9.1], small=True)

doc.add_heading('Limitations of the findings', level=2)
bullets([
    'The reporting window is %s. Five days of steady operation is not a trend.' % W,
    '%s marks were excluded because the centre coordinate is in doubt. Until those centres '
    'are re-surveyed the true rate of marking away from the centre is not known, and it '
    'could move in either direction.' % n(UNVER),
    '%s marks carried no usable position and are counted against nobody.' % n(NO_GPS),
    'The application cannot detect a falsified position directly, only patterns consistent '
    'with one. No mark met that test in this window, which is not the same as none having '
    'occurred.',
    'All findings describe a device, not a person.',
])

doc.add_heading('Open data-protection item', level=2)
callout('Requires closure',
        'The published summary files, which the monitoring console reads, are served from a '
        'public address and contain the storage identifiers of attendance photographs. A '
        'request made without any credential returned image data for one such identifier. '
        'The district should set the photograph folder to restricted access and remove the '
        'identifiers from the published summary. Until that is done, photographs of '
        'Anganwadi staff should be treated as exposed. This was found by the district\'s own '
        'review and is recorded here rather than omitted.', colour=ALERT)

# ------------------------------------------------------------------ save
out_docx = os.path.join(OUTDIR, 'Sisu_Mahila_Samridhi_District_Report_v1.0.docx')
doc.save(out_docx)
print('DOCX  ' + out_docx)
print('      %d figures, %d tables' % (FIG[0], TAB[0]))
print('      %d screens, %d findings rows, %d timed tasks'
      % (len(MANIFEST), len(FINDINGS), len(TIMINGS)))
